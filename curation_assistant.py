"""
Synopsis-Curation-Assistant Tool

- Summarize & Q&A strictly from uploaded files (.pdf, .txt, .doc/.docx, .xlsx)
- Extract figures/images from PDFs, and add to the RAG index
- Upload cBio-style tables (mutations/CNA/SV/clinical) and query gene frequencies
- Query multiple genes (e.g., TP53, EGFR, KRAS) and compute % = (# profiled samples for gene / total profiled samples) * 100

Quickstart
----------
1) Python 3.9+  (3.10+ recommended)
2) pip install -r requirements.txt
3) export OPENAI_API_KEY=sk-...
4) streamlit run new_app.py
"""

from __future__ import annotations
"""
Synopsis-Curation-Assistant Tool
-----------------------------
- Summarize & Q&A strictly from uploaded files (.pdf, .txt, .doc/.docx, .xlsx)
- Extract figures/images from PDFs, and add to the RAG index
- Upload cBio-style tables (mutations/CNA/SV/clinical) and query gene frequencies
- Query multiple genes (e.g., TP53, EGFR, KRAS) and compute % = (# profiled samples for gene / total profiled samples) * 100

"""

# =============================================================================
# SECTION: Compatibility shims (Python 3.13 sqlite, NumPy 2.x alias)
# =============================================================================
import sys
# Provide a modern sqlite module for environments that ship an old stdlib sqlite
# - Requires pysqlite3-binary in requirements. If import fails, we continue
#   without persistence (FAISS fallback will still work).
try:
    import pysqlite3  # type: ignore
    sys.modules["sqlite3"] = sys.modules["pysqlite3"]
    SQLITE_SHIM_OK = True
except Exception:
    SQLITE_SHIM_OK = False

# NumPy 2.x removes np.float_. Some upstream deps still reference it.
try:
    import numpy as np  # type: ignore
    if not hasattr(np, "float_"):
        np.float_ = np.float64  # noqa: NPY201
    NP_SHIM_OK = True
except Exception:
    NP_SHIM_OK = False

# =============================================================================
# SECTION: Standard libs & third-party imports
# =============================================================================
import io
import re
import tempfile
from pathlib import Path
from typing import List, Tuple, Optional

import streamlit as st
import pandas as pd
from PIL import Image

# LangChain + vectorstores
# Try new "langchain_chroma" first; fall back to community path if not installed.
Chroma = None
CLIENT = None
chroma_source = None
client_error = None

try:
    from langchain_chroma import Chroma as _Chroma  # Newer LC integration
    Chroma = _Chroma
    chroma_source = "langchain_chroma"
except Exception:
    try:
        from langchain_community.vectorstores import Chroma as _Chroma  # Legacy LC integration
        Chroma = _Chroma
        chroma_source = "langchain_community"
    except Exception:
        chroma_source = None

# New Chroma Client API (migration-safe). If it fails, we can still run FAISS.
try:
    import chromadb  # noqa: F401
    from chromadb import PersistentClient  # New API
    CLIENT = PersistentClient(path="./chroma_index")
except Exception as e:
    client_error = e

# FAISS fallback (always available to avoid hard failure)
from langchain_community.vectorstores import FAISS

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# Readers
from langchain_community.document_loaders import PyPDFLoader
import docx  # python-docx
import mammoth  # .docx fallback

# =============================================================================
# SECTION: Streamlit setup & constants
# =============================================================================
st.set_page_config(page_title="🧬 Curation Assistant", layout="wide")

# Tuning knobs
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150
TOP_K = 5
MMR_LAMBDA = 0.5

# Persistence & model params
PERSIST_DIR = "./chroma_index"
COLLECTION = "curation_assistant"
EMBED_MODEL = "text-embedding-3-small"
LLM_MODEL = "gpt-4o-mini"

# UI Styles
st.markdown(
    """
<style>
.block-container{padding-top:1.5rem;}
.hero {
  background: linear-gradient(135deg,#6E8EF5 0%, #A777E3 50%, #F08BA9 100%);
  color: white; padding: 18px 20px; border-radius: 16px; margin-bottom: 16px;
  box-shadow: 0 6px 24px rgba(0,0,0,0.15);
}
.tag {background: rgba(255,255,255,0.18); padding: 2px 10px; border-radius: 999px; margin-left: 8px; font-size: 0.85rem;}
.stExpander, .stTabs [data-baseweb="tab"] {border-radius: 12px;}
.stButton>button {border-radius: 999px; padding: 0.5rem 1rem; font-weight: 600;}
</style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="hero">
  <h2 style="margin:0">🧬 Curation Assistant</h2>
  <div style="margin-top:6px">
    Answers grounded strictly in your uploaded files.
    #<span class="tag">Chroma</span><span class="tag">MMR</span><span class="tag">OCR</span><span class="tag">Frequencies</span>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

# Quick env hint
if not st.session_state.get("OPENAI_KEY_WARNED") and not st.secrets.get("OPENAI_API_KEY", None) and not ("OPENAI_API_KEY" in st.session_state or "OPENAI_API_KEY" in st.session_state):
    if not ("OPENAI_API_KEY" in st.session_state or "OPENAI_API_KEY" in st.secrets or "OPENAI_API_KEY" in st.session_state):
        if not ("OPENAI_API_KEY" in st.session_state or "OPENAI_API_KEY" in st.secrets or "OPENAI_API_KEY" in st.session_state):
            pass  # Avoid noisy warnings in some hosts

# =============================================================================
# SECTION: Diagnostics (left sidebar)
# =============================================================================
#st.sidebar.markdown("### 🔧 Diagnostics")
#st.sidebar.write({
    #"python": sys.version,
    #"SQLITE_SHIM_OK": SQLITE_SHIM_OK,
    #"NP_SHIM_OK": NP_SHIM_OK,
    #"Chroma_imported": bool(Chroma),
    #"chroma_source": chroma_source,
    #"chromadb_client_ok": CLIENT is not None,
    #"chromadb_client_error": repr(client_error) if client_error else None,
#})

# =============================================================================
# SECTION: File readers & utilities
# =============================================================================
def _read_txt(file: io.BytesIO, name: str) -> List[Document]:
    """Read plain text-like inputs (txt/md/csv treated as text)."""
    text = file.read().decode("utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source": name})]

def _read_pdf_to_docs(tmp_path: Path, name: str) -> List[Document]:
    """Load a PDF as LangChain Documents (page-wise)."""
    loader = PyPDFLoader(str(tmp_path))
    docs = loader.load()
    for d in docs:
        d.metadata = {**d.metadata, "source": name}
    return docs

def _extract_pdf_images(tmp_path: Path, name: str) -> Tuple[List[Document], List[Image.Image]]:
    """
    Extract embedded images from PDF pages.
    - OCR image regions if pytesseract is available
    - Return both OCR text (as docs) and PIL images for display
    """
    from pypdf import PdfReader
    pil_images: List[Image.Image] = []
    ocr_docs: List[Document] = []
    try:
        reader = PdfReader(str(tmp_path))
        for page_idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            for im in getattr(page, "images", []):
                try:
                    img_bytes = im.data
                    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    pil_images.append(img)
                    ocr_text = ""
                    if "pytesseract" in sys.modules:
                        import pytesseract  # lazy import
                        try:
                            ocr_text = pytesseract.image_to_string(img)
                        except Exception:
                            ocr_text = ""
                    caption_text = "\n".join(
                        [ln for ln in page_text.splitlines() if re.match(r"\s*Fig(ure)?\b", ln, re.I)]
                    )
                    combined = "\n".join([s for s in [caption_text.strip(), ocr_text.strip()] if s])
                    if combined:
                        ocr_docs.append(
                            Document(
                                page_content=f"[FIGURE OCR p.{page_idx}]\n{combined}",
                                metadata={"source": name, "page": page_idx, "type": "figure_ocr"},
                            )
                        )
                except Exception:
                    continue
    except Exception:
        pass
    return ocr_docs, pil_images

def _read_docx(file: io.BytesIO, name: str) -> List[Document]:
    """Read .docx via python-docx; fallback to mammoth raw text."""
    try:
        f = io.BytesIO(file.read())
        doc = docx.Document(f)
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join([p for p in paragraphs if p and p.strip()])
        return [Document(page_content=text, metadata={"source": name})]
    except Exception:
        file.seek(0)
        result = mammoth.extract_raw_text(file)
        text = result.value or ""
        return [Document(page_content=text, metadata={"source": name})]

def _read_xlsx(file: io.BytesIO, name: str) -> List[Document]:
    """Flatten Excel sheets into tab-delimited text documents."""
    docs: List[Document] = []
    try:
        xl = pd.ExcelFile(file)
        for sheet in xl.sheet_names:
            df = xl.parse(sheet).astype(str)
            text_rows = ["\t".join(df.columns.astype(str))]
            text_rows += ["\t".join(map(str, row)) for row in df.itertuples(index=False, name=None)]
            docs.append(Document(page_content="\n".join(text_rows), metadata={"source": name, "sheet": sheet}))
        return docs
    except Exception as e:
        st.error(f"Failed to parse Excel '{name}': {e}")
        return []

def load_files_to_documents(uploaded_files) -> Tuple[List[Document], List[Image.Image]]:
    """Route files by extension and collect both text docs and figures."""
    all_docs: List[Document] = []
    all_figs: List[Image.Image] = []
    for uf in uploaded_files:
        suffix = Path(uf.name).suffix.lower()
        if suffix in [".txt", ".md", ".csv"]:
            all_docs.extend(_read_txt(uf, uf.name))
        elif suffix == ".pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uf.read()); tmp.flush()
                docs = _read_pdf_to_docs(Path(tmp.name), uf.name)
                ocr_docs, figs = _extract_pdf_images(Path(tmp.name), uf.name)
            all_docs.extend(docs); all_docs.extend(ocr_docs); all_figs.extend(figs)
        elif suffix == ".docx":
            all_docs.extend(_read_docx(uf, uf.name))
        elif suffix in [".xlsx", ".xlsm", ".xls"]:
            all_docs.extend(_read_xlsx(uf, uf.name))
        else:
            st.warning(f"Unsupported file type: {uf.name}")
    return all_docs, all_figs

# =============================================================================
# SECTION: Robust table readers (for gene frequency tab)
# =============================================================================
def _rewind(file_obj):
    try: file_obj.seek(0)
    except Exception: pass

def _read_csv_any(file_obj):
    _rewind(file_obj)
    try:
        return pd.read_csv(file_obj, sep=None, engine="python", dtype=str,
                           na_filter=True, low_memory=False, on_bad_lines="skip")
    except Exception:
        _rewind(file_obj)
        return pd.read_csv(file_obj, sep="\t", dtype=str,
                           na_filter=True, low_memory=False, on_bad_lines="skip")

def _read_excel_any(file_obj):
    _rewind(file_obj)
    try:
        buf = io.BytesIO(file_obj.read())
        xls = pd.ExcelFile(buf, engine=None)
        frames = [xls.parse(s, dtype=str) for s in xls.sheet_names]
        return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()
    except Exception as e:
        st.error(f"Excel read failed: {e}")
        return pd.DataFrame()

def _read_any_table(file_obj):
    suffix = Path(file_obj.name).suffix.lower()
    if suffix in {".csv", ".tsv", ".txt"}:
        return _read_csv_any(file_obj)
    elif suffix in {".xlsx", ".xls"}:
        return _read_excel_any(file_obj)
    else:
        st.warning(f"Unsupported supplement file type: {file_obj.name}")
        return pd.DataFrame()

# =============================================================================
# SECTION: Gene frequency helpers
# =============================================================================
def _standardize_cols(cols):
    return [str(c).strip().lower().replace(" ", "_") for c in cols]

def _find_col(cols, candidates):
    for c in cols:
        for cand in candidates:
            if c == cand or c.endswith(f"_{cand}") or cand in c:
                return c
    return None

def infer_total_samples(named_frames, sample_cands=None) -> int:
    """Infer unique sample count across all uploaded supplemental tables."""
    if sample_cands is None:
        sample_cands = [
            "tumor_sample_barcode", "sample_id", "sample", "biosample",
            "patient", "subject", "case_id", "case", "participant_id"
        ]
    uniq = set()
    for _, raw_df in named_frames:
        if raw_df is None or raw_df.empty:
            continue
        df = raw_df.copy()
        df.columns = _standardize_cols(df.columns)
        col = _find_col(df.columns, sample_cands)
        if col and col in df.columns:
            vals = df[col].dropna().astype(str).str.strip()
            uniq.update(vals.tolist())
    return len(uniq)

def compute_gene_frequencies(df: pd.DataFrame, total_samples: Optional[int] = None):
    """
    Accept either long format: (gene, sample_id)
            or aggregated:     (gene, count)
    Return frequency per gene with denominator set to either provided total_samples
    or unique sample count derived from input.
    """
    if df is None or df.empty:
        return pd.DataFrame(columns=["gene", "n_samples", "percentage"]), 0

    df = df.copy()
    df.columns = _standardize_cols(df.columns)

    gene_col   = _find_col(df.columns, ["gene", "symbol", "gene_symbol", "hgnc", "ensembl", "gene_id", "hugo_symbol"])
    sample_col = _find_col(df.columns, [
        "sample", "sample_id", "tumor_sample_barcode", "biosample",
        "patient", "subject", "case_id", "case", "participant_id"
    ])
    count_col  = _find_col(df.columns, ["count", "n", "num", "samples"])

    if not gene_col:
        raise ValueError("Could not detect a gene column (e.g., 'gene', 'symbol', 'Hugo_Symbol').")

    # Long format: gene + sample_id -> unique sample count per gene
    if sample_col and (count_col is None or count_col not in df.columns):
        sub = df[[gene_col, sample_col]].dropna()
        if sub.empty:
            return pd.DataFrame(columns=["gene", "n_samples", "percentage"]), 0
        sub[gene_col] = sub[gene_col].astype(str)
        sub[sample_col] = sub[sample_col].astype(str)
        grp = sub.groupby(gene_col)[sample_col].nunique().reset_index(name="n_samples")
        denom = int(total_samples) if total_samples else int(sub[sample_col].nunique())

    # Aggregated: gene + count
    elif count_col in df.columns:
        sub = df[[gene_col, count_col]].dropna()
        if sub.empty:
            return pd.DataFrame(columns=["gene", "n_samples", "percentage"]), 0
        sub[gene_col] = sub[gene_col].astype(str)
        sub[count_col] = pd.to_numeric(sub[count_col], errors="coerce").fillna(0).astype(int)
        grp = sub.groupby(gene_col)[count_col].sum().reset_index(name="n_samples")
        denom = int(total_samples) if total_samples else int(grp["n_samples"].sum())

    else:
        raise ValueError(
            "Could not detect suitable columns. Expected either (gene + sample_id) or (gene + count). "
            f"Columns seen: {list(df.columns)[:12]} ..."
        )

    grp = grp.sort_values("n_samples", ascending=False)
    denom = max(int(denom), 1)
    grp["percentage"] = (grp["n_samples"] / denom) * 100.0
    grp["percentage"] = grp["percentage"].round(6)
    if grp.columns[0] != "gene":
        grp = grp.rename(columns={grp.columns[0]: "gene"})
    return grp, denom

# =============================================================================
# SECTION: Index builders (Chroma new client API + FAISS fallback)
# =============================================================================
def _split_docs(docs: List[Document]):
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    splits = splitter.split_documents(docs)
    if len(splits) == 0:
        raise ValueError("No text extracted from the uploaded files.")
    return splits

def build_index(docs: List[Document]):
    """
    Try Chroma (persistent). If the Chroma client/vectorstore cannot be created
    (e.g., due to sqlite version or missing chroma), fall back to FAISS in-memory.
    """
    splits = _split_docs(docs)
    embeddings = OpenAIEmbeddings(model=EMBED_MODEL)

    # Attempt Chroma first (requires both Chroma class and a working PersistentClient)
    if Chroma is not None and CLIENT is not None:
        try:
            vs = Chroma.from_documents(
                splits,
                embeddings,
                client=CLIENT,                # New API: pass the PersistentClient
                collection_name=COLLECTION,
            )
            try:
                vs.persist()  # safe-guarded (may be a no-op with some builds)
            except Exception:
                pass
            backend = "chroma"
            return vs, embeddings, splits, backend
        except Exception as e:
            st.sidebar.error(f"Chroma init failed; falling back to FAISS: {e}")

    # FAISS: fast, reliable, in-memory (no persistence)
    vs = FAISS.from_documents(splits, embeddings)
    backend = "faiss"
    return vs, embeddings, splits, backend

def load_existing_index():
    """
    Load an existing Chroma collection if available; FAISS has no persistence
    so we create a new embedding function and return None for the vectorstore.
    """
    embeddings = OpenAIEmbeddings(model=EMBED_MODEL)
    if Chroma is not None and CLIENT is not None:
        try:
            vs = Chroma(
                client=CLIENT,
                collection_name=COLLECTION,
                embedding_function=embeddings,
            )
            return vs, embeddings, "chroma"
        except Exception as e:
            st.sidebar.error(f"Chroma load failed; new ingestion needed: {e}")

    # No persisted FAISS — user must (re)ingest
    return None, embeddings, "faiss"

# =============================================================================
# SECTION: Prompt templates
# =============================================================================
SUMMARY_SYSTEM = (
    "You are a meticulous scientific analyst. Generate a clear, concise, and strictly factual summary of the provided "
    "content. Do not speculate or add external knowledge. If tables exist, highlight the most important columns and "
    "notable values. If figures are present, briefly describe key features they illustrate. Keep the summary grounded."
)
SUMMARY_USER = "Summarize the following corpus in <= 20 bullet points. If multiple files are present, group by source name.\n\n{context}"
summary_prompt = ChatPromptTemplate.from_messages([("system", SUMMARY_SYSTEM), ("human", SUMMARY_USER)])

QA_SYSTEM = (
    "You are a bioinformatics software engineer. Answer ONLY using the given context. "
    "If the answer isn't in the context, say 'I don't know from the provided files.' "
    "Provide brief citations like (source, optional sheet/page)."
)
QA_USER = "Question: {question}\n\nUse the context to answer concisely in 1-15 sentences.\n\nContext:\n{context}"
qa_prompt = ChatPromptTemplate.from_messages([("system", QA_SYSTEM), ("human", QA_USER)])

# =============================================================================
# SECTION: Streamlit session state
# =============================================================================
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "all_docs" not in st.session_state:
    st.session_state.all_docs = []
if "all_figs" not in st.session_state:
    st.session_state.all_figs = []
if "supp_named_frames" not in st.session_state:
    st.session_state.supp_named_frames = []
if "backend" not in st.session_state:
    st.session_state.backend = None

# =============================================================================
# SECTION: Sidebar (uploads & actions)
# =============================================================================
with st.sidebar:
    st.subheader("Upload files")
    uploads = st.file_uploader(
        "Drop PDF/TXT/DOC/DOCX/XLSX files",
        type=["pdf", "txt", "md", "csv", "doc", "docx", "xlsx", "xlsm", "xls"],
        accept_multiple_files=True,
    )
    colb1, colb2 = st.columns(2)
    with colb1:
        build_btn = st.button("Ingest Files", type="primary")
    with colb2:
        load_btn = st.button("Load Existing Index")

# =============================================================================
# SECTION: Index actions
# =============================================================================
if build_btn:
    if not uploads:
        st.warning("Please upload at least one file.")
    else:
        with st.spinner("Parsing files, extracting figures & embedding…"):
            docs, figs = load_files_to_documents(uploads)
            if len(docs) == 0:
                st.error("No parsable content found in the uploaded files.")
            else:
                try:
                    vs, _, splits, backend = build_index(docs)
                    st.session_state.vectorstore = vs
                    st.session_state.all_docs = splits
                    st.session_state.all_figs = figs
                    st.session_state.backend = backend
                    st.success(f"Indexed {len(splits)} chunks with backend: {backend.upper()}. Figures: {len(figs)}.")
                except Exception as e:
                    st.exception(e)

if load_btn:
    with st.spinner("Loading existing index…"):
        try:
            vs, _, backend = load_existing_index()
            if vs is not None:
                st.session_state.vectorstore = vs
                st.session_state.backend = backend
                st.info(f"Loaded existing collection with backend: {backend.upper()}.")
            else:
                st.warning("No persisted index found; please ingest files.")
        except Exception as e:
            st.exception(e)

# =============================================================================
# SECTION: Tabs (Summary, QA, Figures, Frequencies)
# =============================================================================
summary_tab, qa_tab, figs_tab, freq_tab = st.tabs(
    ["📝 Summary", "❓ Q&A", "🖼️ Figures", "📊 Gene Frequencies"]
)

# ----- Summary -----
with summary_tab:
    st.write("Generate a concise overview of the ingested corpus.")
    if st.session_state.vectorstore is None and not st.session_state.all_docs:
        st.info("Build or load an index to enable summarization.")
    else:
        if st.button("Generate Summary", key="summ_btn"):
            with st.spinner("Summarizing…"):
                subset = list(st.session_state.all_docs)[:25]  # small cap for faster demo
                if not subset:
                    st.warning("No chunks available to summarize. Try rebuilding the index.")
                else:
                    llm = ChatOpenAI(model=LLM_MODEL, temperature=0)
                    chain = summary_prompt | llm | StrOutputParser()
                    context_text = "\n\n".join([
                        f"[src={d.metadata.get('source','-')} sheet={d.metadata.get('sheet','-')} page={d.metadata.get('page','-')}]\n{d.page_content}"
                        for d in subset
                    ])
                    summary = chain.invoke({"context": context_text})
                    st.markdown(summary)

# ----- Q&A -----
with qa_tab:
    st.write("Ask questions strictly from your files. If it’s not in them, the app will say it doesn’t know.")
    if st.session_state.vectorstore is None:
        st.info("Build or load an index to enable Q&A.")
    else:
        q = st.text_input("Your question")
        if st.button("Ask", type="primary") and q.strip():
            with st.spinner("Retrieving evidence & answering…"):
                retriever = st.session_state.vectorstore.as_retriever(
                    search_type="mmr",
                    search_kwargs={"k": TOP_K, "lambda_mult": MMR_LAMBDA},
                )
                docs = retriever.get_relevant_documents(q)
                if not docs:
                    st.warning("No sufficiently relevant context found. I don't know from the provided files.")
                else:
                    context_text = "\n\n".join([
                        f"[src={d.metadata.get('source','-')} sheet={d.metadata.get('sheet','-')} page={d.metadata.get('page','-')}]\n{d.page_content}"
                        for d in docs
                    ])
                    llm = ChatOpenAI(model=LLM_MODEL, temperature=0)
                    chain = qa_prompt | llm | StrOutputParser()
                    answer = chain.invoke({"question": q, "context": context_text})
                    st.markdown("**Answer**")
                    st.write(answer)

                    st.markdown("**Cited Chunks**")
                    for i, d in enumerate(docs, 1):
                        src = d.metadata.get("source")
                        sheet = d.metadata.get("sheet")
                        page = d.metadata.get("page")
                        with st.expander(f"{i}. {src} — sheet: {sheet or '-'} — page: {page or '-'}"):
                            st.code(d.page_content[:2000])

# ----- Figures -----
with figs_tab:
    st.write("Extracted figures from PDFs (OCR text included in the index when available).")
    figs = st.session_state.all_figs
    if not figs:
        st.info("No figures extracted yet — upload PDFs and rebuild the index.")
    else:
        cols = st.columns(3)
        for i, img in enumerate(figs):
            with cols[i % 3]:
                st.image(img, caption=f"Figure {i+1}", use_column_width=True)

# ----- Gene Frequencies -----
with freq_tab:
    st.write("Upload **cBio-style tables** (mutations/CNA/SV/clinical). Then query gene frequencies below.")

    # Upload supplemental tables
    supps = st.file_uploader(
        "Upload supplementary tables (CSV/TSV/TXT/XLSX)",
        type=["csv", "tsv", "txt", "xlsx", "xls"],
        accept_multiple_files=True,
        key="supp_upload",
    )

    named_frames = []
    if supps:
        for f in supps:
            df = _read_any_table(f)
            if not df.empty:
                named_frames.append((f.name, df))
                st.write(f"✔️ Loaded **{f.name}**  shape={df.shape}")
            else:
                st.warning(f"Skipping empty/unreadable file: {f.name}")
        if named_frames:
            st.session_state["supp_named_frames"] = named_frames

    st.markdown("---")
    st.subheader("🔎 Query Gene Frequency (mutations/CNA/SV) — multiple genes")

    if "supp_named_frames" not in st.session_state or not st.session_state["supp_named_frames"]:
        st.info("Upload files above, then enter genes here.")
    else:
        named_frames = st.session_state["supp_named_frames"]

        genes_input = st.text_input("Gene symbols (comma-separated, e.g., TP53, EGFR, KRAS)")
        use_case_insensitive = st.checkbox("Case-insensitive match", value=True)

        auto_denom = infer_total_samples(named_frames)
        denom_override = st.number_input(
            f"Total number of profiled samples (0 = auto; auto currently infers {auto_denom})",
            min_value=0, step=1, value=0
        )
        denominator = int(denom_override) if denom_override > 0 else int(auto_denom)

        SAMPLE_COL_CANDS = [
            "tumor_sample_barcode", "sample_id", "sample", "biosample",
            "patient", "subject", "case_id", "case", "participant_id"
        ]

        def _norm(df: pd.DataFrame) -> pd.DataFrame:
            d = df.copy(); d.columns = _standardize_cols(d.columns); return d

        def _find_sample_col(cols):
            return _find_col(cols, SAMPLE_COL_CANDS)

        def compute_numerator_for_gene(gene_query: str):
            """Count distinct samples per gene across all uploaded supplemental tables."""
            gene_upper = gene_query.strip().upper()
            numerator_samples = set()
            add_rows_without_ids = 0
            per_file_counts = []

            for fname, raw_df in named_frames:
                if raw_df is None or raw_df.empty:
                    continue
                df = _norm(raw_df)
                cols = list(df.columns)

                hugo_col = _find_col(cols, ["hugo_symbol"])
                site1_col = _find_col(cols, ["site1_hugo_symbol"])
                site2_col = _find_col(cols, ["site2_hugo_symbol"])
                sample_col = _find_sample_col(cols)

                file_added = 0

                # Mutation/CNA long tables
                if hugo_col:
                    mask = (
                        df[hugo_col].astype(str).str.strip().str.upper() == gene_upper
                        if use_case_insensitive else
                        df[hugo_col].astype(str) == gene_query
                    )
                    sub = df.loc[mask]
                    if not sub.empty:
                        if sample_col in sub.columns:
                            ids = sub[sample_col].dropna().astype(str).str.strip()
                            ids = set(ids.tolist())
                            numerator_samples |= ids
                            file_added = len(ids)
                        else:
                            # Wide tables w/out explicit sample IDs: approximate by non-null sample-like cols
                            non_id_cols = [c for c in sub.columns if c != hugo_col]
                            meta_like = {"entrez_gene_id", "chromosome", "cytoband"}
                            non_id_cols = [c for c in non_id_cols if c not in meta_like]
                            if len(non_id_cols) > 5:
                                nn = int(sub[non_id_cols].notna().sum(axis=1).max())
                                add_rows_without_ids += nn
                                file_added = nn
                            else:
                                cnt = int(len(sub))
                                add_rows_without_ids += cnt
                                file_added = cnt

                # Structural variants
                elif site1_col or site2_col:
                    m1 = df[site1_col].astype(str).str.strip().str.upper() == gene_upper if site1_col else False
                    m2 = df[site2_col].astype(str).str.strip().str.upper() == gene_upper if site2_col else False
                    sub = df.loc[m1 | m2]
                    if not sub.empty:
                        if sample_col in sub.columns:
                            ids = sub[sample_col].dropna().astype(str).str.strip()
                            ids = set(ids.tolist())
                            numerator_samples |= ids
                            file_added = len(ids)
                        else:
                            cnt = int(len(sub))
                            add_rows_without_ids += cnt
                            file_added = cnt

                per_file_counts.append((fname, file_added))

            numerator = len(numerator_samples) + add_rows_without_ids
            return int(numerator), per_file_counts

        if st.button("Compute Gene Frequency") and genes_input.strip():
            # Parse comma-separated genes, dedupe preserving order
            raw_list = [g.strip() for g in genes_input.split(",") if g.strip()]
            seen = set(); genes = []
            for g in raw_list:
                key = g.upper() if use_case_insensitive else g
                if key not in seen:
                    seen.add(key); genes.append(g)

            if denominator == 0:
                st.warning("Denominator could not be inferred. Please set a value above.")
            else:
                results = []
                per_gene_breakdown = {}
                for g in genes:
                    num, breakdown = compute_numerator_for_gene(g)
                    per_gene_breakdown[g] = breakdown
                    pct = (num / denominator * 100.0)
                    results.append({
                        "gene": g,
                        "n_samples": int(num),
                        "denominator": int(denominator),
                        "percentage": round(pct, 6),
                    })

                st.markdown("**Per-gene results**")
                st.dataframe(pd.DataFrame(results))

                with st.expander("Per-file breakdown (by gene)"):
                    for g in genes:
                        st.markdown(f"**{g}**")
                        rows = per_gene_breakdown.get(g, [])
                        if rows:
                            st.table(pd.DataFrame(rows, columns=["file", "# counted"]))
                        else:
                            st.write("_no rows counted_")
